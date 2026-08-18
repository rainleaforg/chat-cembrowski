#type: ignore
"""Ingestion module for miscellaneous context documents (txt, docx, code files).

Scans data/docs/ and creates Document objects with structured extracted text.
Run with: uv run -m chat_cembrowski.data.doc_ingestion
"""

import hashlib
import logging
import re
import uuid
from pathlib import Path
from typing import Optional

import yaml

from .models import Document
from .serialization import save_document, load_documents_from_json

logger = logging.getLogger(__name__)

DOCS_DIR = Path(__file__).resolve().parents[3] / "data" / "docs"
DOC_JSON_DIR = Path(__file__).resolve().parents[3] / "data" / "doc_json"

# Developer-authored site/internal knowledge markdown (see scripts/ingest_knowledge.py)
# — a separate source directory from data/docs/, ingested straight to Qdrant
# rather than through vectordb.py's paper-centric __main__ loop.
KNOWLEDGE_DIR = Path(__file__).resolve().parents[3] / "data" / "knowledge"
KNOWLEDGE_JSON_DIR = Path(__file__).resolve().parents[3] / "data" / "knowledge_json"

# Fixed namespace for deriving a stable Document ID from its filename, so an
# edited file maps back to the same record (and therefore the same chunk IDs)
# instead of being re-registered as a new document with orphaned old chunks.
DOC_ID_NAMESPACE = uuid.uuid5(uuid.NAMESPACE_DNS, "chat_cembrowski.documents")

TEXT_EXTENSIONS = {".txt", ".md"}
DOCX_EXTENSIONS = {".docx"}
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx",
    ".c", ".cpp", ".h", ".hpp",
    ".java", ".r", ".m", ".sh", ".bash",
    ".sql", ".yaml", ".yml", ".toml", ".ini", ".cfg",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCX_EXTENSIONS | CODE_EXTENSIONS


def doc_id_for(source_file: str) -> str:
    """Stable Document ID derived from the filename."""
    return str(uuid.uuid5(DOC_ID_NAMESPACE, source_file))


def content_hash(text: str) -> str:
    """
    sha256 of a document's extracted text.

    Hashing the *extracted* text rather than the raw bytes means a .docx
    re-saved with no content change (which rewrites its zip container and
    changes the file bytes) does not trigger a pointless re-embed.
    """
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _extract_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(file_path: Path) -> str:
    """Extract text from a .docx file as structured markdown.

    Headings → # / ## / ###, list paragraphs → - bullets, tables → markdown tables.
    This structure improves chunking by giving the splitter clear semantic boundaries.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))

    # Build element → object lookups for O(1) access during body traversal.
    para_by_elem = {p._element: p for p in doc.paragraphs}
    table_by_elem = {t._element: t for t in doc.tables}

    lines: list[str] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]  # strip XML namespace

        if tag == "p":
            para = para_by_elem.get(block)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 4
                lines.append(f"{'#' * level} {text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            tbl = table_by_elem.get(block)
            if tbl is None:
                continue

            lines.append("")
            for i, row in enumerate(tbl.rows):
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

    return "\n".join(lines)


def _extract_code(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="replace")


_FRONT_MATTER_RE = re.compile(r"\A---[ \t]*\n(.*?\n)---[ \t]*\n?", re.DOTALL)


def _parse_front_matter(text: str) -> tuple[dict, str]:
    """
    Split a leading YAML front-matter block (title/kind/site_path) off a
    markdown file's text. Returns ({}, text) unchanged when there is none.

    The front-matter is stripped from the returned body so it never reaches
    the model as prose — kind/site_path are metadata, not content.
    """
    match = _FRONT_MATTER_RE.match(text)
    if not match:
        return {}, text
    try:
        metadata = yaml.safe_load(match.group(1)) or {}
    except yaml.YAMLError as e:
        logger.warning(f"Malformed front-matter, ignoring it: {e}")
        return {}, text
    # safe_load happily returns a scalar or a list for a well-formed block
    # that just isn't a mapping ("kind: site" without the key, a bare list).
    # Callers read metadata with .get(), so anything else is malformed here.
    if not isinstance(metadata, dict):
        logger.warning(
            "Front-matter is %s, not a mapping — ignoring it.",
            type(metadata).__name__,
        )
        return {}, text
    return metadata, text[match.end():]


def ingest_local_docs(
    docs_dir: Path = DOCS_DIR,
    doc_json_dir: Path = DOC_JSON_DIR,
) -> list[Document]:
    """Create Document objects for files in docs_dir not yet registered in doc_json_dir.

    Supported types: .txt, .md, .docx, and common code extensions.
    Idempotent — skips files whose source_file is already in doc_json_dir.

    Args:
        docs_dir: Source directory (default: data/docs)
        doc_json_dir: JSON output directory (default: data/doc_json)

    Returns:
        List of newly created Document objects.
    """
    docs_dir.mkdir(parents=True, exist_ok=True)
    doc_json_dir.mkdir(parents=True, exist_ok=True)

    existing = {d.source_file: d for d in load_documents_from_json(doc_json_dir)}

    new_docs: list[Document] = []
    updated_docs: list[Document] = []

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            logger.debug(f"Skipping unsupported file: {file_path.name}")
            continue

        ext = file_path.suffix.lower()

        try:
            if ext in DOCX_EXTENSIONS:
                text = _extract_docx(file_path)
                file_type = "docx"
            elif ext in TEXT_EXTENSIONS:
                text = _extract_txt(file_path)
                file_type = ext.lstrip(".")
            else:
                text = _extract_code(file_path)
                file_type = ext.lstrip(".")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name}: {e}")
            continue

        front_matter: dict = {}
        if ext == ".md":
            front_matter, text = _parse_front_matter(text)

        title = front_matter.get("title") or file_path.stem
        kind = front_matter.get("kind", "document")
        site_path = front_matter.get("site_path")

        if kind == "site" and not site_path:
            logger.error(
                f"'{file_path.name}': kind is 'site' but front-matter has no "
                "site_path — skipping. A site chunk with no link can't be cited."
            )
            continue

        digest = content_hash(text)
        prior = existing.get(file_path.name)

        if prior is not None:
            # kind/site_path affect only the chunk payload, not the embedded
            # text, so they're excluded from content_hash — but a front-matter-
            # only edit still has to trigger re-ingestion, or the old payload
            # values would never get refreshed.
            metadata_changed = (
                prior.title != title or prior.kind != kind or prior.site_path != site_path
            )
            if prior.content_hash == digest and not metadata_changed:
                logger.info(f"Unchanged, skipping: {file_path.name}")
                continue

            # Edited in place. Keep the ID (so chunk IDs stay stable) and clear
            # processed so vectordb re-indexes: it deletes this doc_id's points
            # before upserting, which also clears chunks an edit removed.
            if not prior.content_hash:
                reason = "no stored hash"
            elif prior.content_hash != digest:
                reason = "content changed"
            else:
                reason = "front-matter changed"
            logger.info(f"Re-ingesting '{file_path.name}' ({reason}).")
            prior.text = text
            prior.file_type = file_type
            prior.title = title
            prior.kind = kind
            prior.site_path = site_path
            prior.content_hash = digest
            prior.processed = False
            save_document(prior, doc_json_dir)
            updated_docs.append(prior)
            continue

        doc = Document(
            id=doc_id_for(file_path.name),
            title=title,
            source_file=file_path.name,
            file_type=file_type,
            text=text,
            content_hash=digest,
            processed=False,
            kind=kind,
            site_path=site_path,
        )
        save_document(doc, doc_json_dir)
        new_docs.append(doc)
        logger.info(
            f"Created Document: '{doc.title}' ({doc.file_type}, kind={doc.kind}, {len(doc.text):,} chars)"
        )

    logger.info(
        f"Document ingestion complete. New: {len(new_docs)}, updated: {len(updated_docs)}."
    )
    return new_docs + updated_docs


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    docs = ingest_local_docs()
    for doc in docs:
        print(f"- {doc.title} ({doc.file_type}, {len(doc.text):,} chars)")
